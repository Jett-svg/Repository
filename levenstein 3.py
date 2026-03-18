import os
from docx import Document
from fuzzywuzzy import fuzz
from fuzzywuzzy import process
import Levenshtein

def read_text_from_docx():
    try:
        doc = Document(filepath)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                full_text.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    full_text.append(' '.join(row_text))

        return '\n'.join(full_text)

    except Exception as e:
        print(f"Ошибка при чтении файла {filepath}: {e}")
        return ""


def compare_strings_fuzz(str1, str2, method='ratio'):
    methods = {
        'ratio': fuzz.ratio,
        'partial_ratio': fuzz.partial_ratio,
        'token_sort_ratio': fuzz.token_sort_ratio,
        'token_set_ratio': fuzz.token_set_ratio,
        'wratio': fuzz.WRatio,
    }

    if method not in methods:
        raise ValueError(f"Метод {method} не поддерживается")

    return methods[method](str1, str2)


def calculate_levenshtein_distance(str1, str2):
    return Levenshtein.distance(str1, str2)

def calculate_similarity_ratio(str1, str2):
    return Levenshtein.ratio(str1, str2)


def compare_two_documents(file1, file2, method='wratio'):
    text1 = read_word_file(file1)
    text2 = read_word_file(file2)

    if not text1 or not text2:
        return {"error": "Не удалось прочитать один или оба файла"}

    results = {}

    methods_to_test = ['ratio', 'partial_ratio', 'token_sort_ratio',
                       'token_set_ratio', 'wratio']

    for method in methods_to_test:
        score = compare_strings_fuzz(text1, text2, method)
        results[f'fuzz_{method}'] = score

    results['levenshtein_distance'] = calculate_levenshtein_distance(text1, text2)
    results['levenshtein_ratio'] = calculate_similarity_ratio(text1, text2)

    return results

if __name__ == "__main__":
    result = compare_two_documents("document1.docx", "document2.docx")
    for method, score in result.items():
        print(f"{method}: {score}")


def compare_multiple_documents(file_list, reference_file=None):
    documents = {}
    for filepath in file_list:
        text = read_word_file(filepath)
        if text:
            documents[filepath] = text

    results = {}

    if reference_file:
        ref_text = documents.get(reference_file, read_word_file(reference_file))
        if ref_text:
            for filepath, text in documents.items():
                if filepath != reference_file:
                    score = fuzz.WRatio(ref_text, text)
                    results[filepath] = score
    else:
        file_paths = list(documents.keys())
        for i in range(len(file_paths)):
            for j in range(i + 1, len(file_paths)):
                pair_key = f"{os.path.basename(file_paths[i])} vs {os.path.basename(file_paths[j])}"
                score = fuzz.WRatio(documents[file_paths[i]], documents[file_paths[j]])
                results[pair_key] = score

    return results

